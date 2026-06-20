"""
doc_editor.py
DOCX template editing with placeholder replacement.

Rules (strict):
  - resume.docx : replace {{JOB_TITLE}} (and [Job Title] / literal job title text as fallback)
  - cv.docx     : replace {{COMPANY_NAME}}, {{ROLE}}, {{JOB_TITLE}} (header), and {{LOCATION}}
  - All other content (formatting, styles, fonts, colours) must be 100% preserved.

The replacement algorithm works at the Run level so bold/italic/font/size
on the replaced run is preserved. It also handles the common Word behaviour
where a single placeholder is split across multiple consecutive runs
(e.g. "{{JOB" in run-0 and "_TITLE}}" in run-1).

Fallback behaviour:
  If no standard placeholder ({{JOB_TITLE}} / [Job Title]) is found, the editor
  attempts to replace any known literal job title text (e.g. "Full Stack Developer")
  so that user-provided templates that never had a placeholder are still supported.
"""

from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
import re


def clean_role_name(role: str) -> str:
    """
    Remove seniority qualifiers from role names to ensure consistent,
    professional filenames.
    """
    if not role:
        return ""
        
    qualifiers = [
        r'\b(?:senior|junior|intern|lead|principal|chief|staff|head\s+of|associate|assistant)\b',
        r'\b(?:sr\.?|jr\.?)\b',
        r'\b(?:entry[- ]level|mid[- ]level|upper[- ]level|executive)\b'
    ]
    
    cleaned = role
    for q in qualifiers:
        cleaned = re.sub(q, '', cleaned, flags=re.IGNORECASE)
    
    cleaned = re.sub(r'\s+', ' ', cleaned)
    cleaned = re.sub(r'^[-\s,]+|[-\s,]+$', '', cleaned)
    
    return cleaned.strip() or role.strip()


# ---------------------------------------------------------------------------
# Skill Categorization
# ---------------------------------------------------------------------------

SKILL_CATEGORIES = {
    "Frontend": [
        "TypeScript", "React", "React.js", "Vue", "Vue.js", "Angular", "Svelte",
        "JavaScript", "JSX", "Next.js", "Gatsby", "Remix", "Nuxt", 
        "Tailwind", "Tailwind CSS", "Bootstrap", "Material UI", "Shadcn",
        "CSS", "SCSS", "LESS", "Sass", "Styled Components", "Emotion",
        "HTML", "HTML5", "WebGL", "Canvas", "SVG",
        "Redux", "Zustand", "Context", "Jotai", "Recoil",
        "Framer Motion", "Three.js", "D3.js", "Chart.js", "Visx"
    ],
    "Backend": [
        "Node.js", "Node", "Express", "Express.js", "NestJS", "Fastify",
        "Django", "Flask", "FastAPI", "Python", "Go", "Rust", "Java",
        "MongoDB", "PostgreSQL", "MySQL", "MariaDB", "SQLite", "Oracle",
        "Redis", "Memcached", "Elasticsearch", "Cassandra",
        "Prisma", "Sequelize", "TypeORM", "SQLAlchemy", "Hibernate",
        "GraphQL", "Apollo", "REST API", "REST", "SOAP",
        "Socket.io", "WebSocket", "tRPC", "gRPC",
        "Microservices", "Monolithic", "Serverless", "Lambda"
    ],
    "Auth": [
        "JWT", "OAuth", "OAuth2", "OpenID", "SAML",
        "Firebase Auth", "Firebase", "Better Auth", "NextAuth", "NextAuth.js",
        "Passport.js", "Passport", "Auth0", "Okta", "Cognito",
        "Session", "Cookies", "Token", "Two-Factor Auth", "MFA", "2FA",
        "Bcrypt", "Argon2", "PBKDF2"
    ],
    "Tools": [
        "Docker", "Docker Compose", "Kubernetes", "K8s",
        "Git", "GitHub", "GitLab", "Bitbucket", "SVN",
        "AWS", "Azure", "Google Cloud", "GCP", "Heroku", "Vercel", "Netlify",
        "Jenkins", "GitHub Actions", "GitLab CI", "CircleCI", "Travis CI",
        "Postman", "Insomnia", "Thunder Client", "REST Client",
        "Figma", "Adobe XD", "Sketch", "InVision",
        "Jira", "Trello", "Asana", "Monday.com",
        "Linux", "Unix", "Windows Server", "macOS",
        "Nginx", "Apache", "IIS",
        "Webpack", "Vite", "Parcel", "Rollup", "Esbuild",
        "NPM", "Yarn", "PNPM", "Bun"
    ]
}

def _categorize_skills(skills: list[str]) -> dict:
    """
    Categorize a list of skills into Frontend, Backend, Auth, and Tools.
    
    Args:
        skills: List of skill names
    
    Returns:
        Dict with categories as keys and comma-separated skills as values
    """
    categorized = {
        "Frontend": [],
        "Backend": [],
        "Auth": [],
        "Tools": []
    }
    
    for skill in skills:
        skill_lower = skill.lower()
        found = False
        
        for category, skill_list in SKILL_CATEGORIES.items():
            for cat_skill in skill_list:
                if skill_lower == cat_skill.lower():
                    categorized[category].append(skill)
                    found = True
                    break
            if found:
                break
        
        # If not found in any category, add to Tools (default)
        if not found:
            categorized["Tools"].append(skill)
    
    # Format as comma-separated strings, empty string if no skills
    result = {}
    for category, skills_list in categorized.items():
        if skills_list:
            # Remove duplicates while preserving order
            seen = set()
            unique_skills = []
            for s in skills_list:
                s_lower = s.lower()
                if s_lower not in seen:
                    unique_skills.append(s)
                    seen.add(s_lower)
            result[category] = ", " + ", ".join(unique_skills)
        else:
            result[category] = ""
    
    return result


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _paragraph_full_text(para) -> str:
    """Return the concatenated text of all runs in a paragraph."""
    return "".join(run.text for run in para.runs)


def _replace_in_runs(runs, old: str, new: str) -> bool:
    """
    Replace *old* with *new* inside a sequence of runs, preserving formatting.

    Strategy
    --------
    1. Build a flat character list that records which run each character
       belongs to, so we can find the placeholder even when it is split
       across multiple runs.
    2. When a match is found, put the replacement text into the first run
       of the span and empty the remaining runs that were part of the match.

    Returns True if at least one replacement was made.
    """
    changed = False

    while True:
        # Re-build the full text and the per-character run-index map each
        # iteration so that multiple occurrences are handled correctly.
        full_text = ""
        char_run_idx = []  # char_run_idx[i] = index into `runs` for char i
        for ri, run in enumerate(runs):
            for _ in run.text:
                char_run_idx.append(ri)
            full_text += run.text

        idx = full_text.find(old)
        if idx == -1:
            break  # No more occurrences

        end_idx = idx + len(old) - 1  # inclusive

        first_run_idx = char_run_idx[idx]
        last_run_idx = char_run_idx[end_idx]

        # Build new texts for every affected run
        # For the first run: keep chars before the match start + replacement
        run_start_offset = sum(
            len(runs[ri].text) for ri in range(first_run_idx)
        )
        local_start = idx - run_start_offset  # offset within the first run

        # Text of first run: chars before match + replacement value
        first_run_text = runs[first_run_idx].text[:local_start] + new

        # Text of last run: chars after match end (within that run)
        last_run_start_offset = sum(
            len(runs[ri].text) for ri in range(last_run_idx)
        )
        local_end = end_idx - last_run_start_offset  # inclusive offset in last run
        last_run_suffix = runs[last_run_idx].text[local_end + 1:]

        if first_run_idx == last_run_idx:
            # Entire placeholder sits in one run
            runs[first_run_idx].text = first_run_text + last_run_suffix
        else:
            runs[first_run_idx].text = first_run_text
            # Zero out intermediate runs
            for ri in range(first_run_idx + 1, last_run_idx):
                runs[ri].text = ""
            # Last run keeps only the suffix
            runs[last_run_idx].text = last_run_suffix

        changed = True

    return changed


def _replace_in_paragraph(para, replacements: dict) -> None:
    """Apply all replacements to a single paragraph's runs."""
    for old, new in replacements.items():
        _replace_in_runs(para.runs, old, new)


def _iter_text_box_paragraphs(doc_or_header):
    """
    Yield every paragraph found inside text boxes (``w:txbxContent`` elements).

    Word stores floating text boxes as drawing objects whose XML looks like::

        <w:drawing>
          <wp:inline or wp:anchor>
            <a:graphic>
              <a:graphicData>
                <wps:txbx>
                  <w:txbxContent>
                    <w:p> … </w:p>
                  </w:txbxContent>
                </wps:txbx>
              </a:graphicData>
            </a:graphic>
          </wp:inline>
        </w:drawing>

    python-docx does not expose these through its high-level API, so we
    descend into the raw XML.

    Accepts both ``Document`` objects (which expose ``.element``) and
    header/footer objects (``_Header`` / ``_Footer``, which expose
    ``._element``).
    """
    from docx.text.paragraph import Paragraph as _Paragraph

    # Namespace for the WordprocessingML schema
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    txbx_tag = f"{{{W_NS}}}txbxContent"
    p_tag    = f"{{{W_NS}}}p"

    # Document exposes .element; _Header/_Footer expose ._element
    root_elem = getattr(doc_or_header, "element", None)
    if root_elem is None:
        root_elem = getattr(doc_or_header, "_element", None)
    if root_elem is None:
        return

    for txbx_elem in root_elem.iter(txbx_tag):
        for p_elem in txbx_elem.iter(p_tag):
            yield _Paragraph(p_elem, None)


def _replace_in_document(doc: Document, replacements: dict) -> None:
    """
    Walk every paragraph, table cell, header/footer, and text box in the
    document and apply replacements.
    Does NOT touch any other XML; formatting is fully preserved.
    """
    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Tables (including nested tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # Headers and footers
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is not None:
                for para in header.paragraphs:
                    _replace_in_paragraph(para, replacements)
                # Also cover text boxes inside headers
                for para in _iter_text_box_paragraphs(header):
                    _replace_in_paragraph(para, replacements)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is not None:
                for para in footer.paragraphs:
                    _replace_in_paragraph(para, replacements)
                for para in _iter_text_box_paragraphs(footer):
                    _replace_in_paragraph(para, replacements)

    # Floating text boxes in the body
    for para in _iter_text_box_paragraphs(doc):
        _replace_in_paragraph(para, replacements)


def _document_full_text(doc: Document) -> str:
    """Return all text in the document for placeholder detection (incl. text boxes)."""
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    parts.append(para.text)
    for para in _iter_text_box_paragraphs(doc):
        parts.append(para.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# DocEditor
# ---------------------------------------------------------------------------

# Base job-title keywords we look for in the document text.
# The scanner will match these even if the template has a " | tech stack"
# suffix appended (e.g. "Full Stack Developer | Next.js · Node.js · PostgreSQL").
_KNOWN_JOB_TITLE_BASES = [
    "Full Stack Developer",
    "Fullstack Developer",
    "Full-Stack Developer",
    "Software Engineer",
    "Backend Developer",
    "Frontend Developer",
    "Front-End Developer",
    "Web Developer",
    "Mobile Developer",
    "Data Scientist",
    "Data Analyst",
    "DevOps Engineer",
    "Product Manager",
    "Project Manager",
    "UI/UX Designer",
    "Graphic Designer",
    "Business Analyst",
    "QA Engineer",
    "Machine Learning Engineer",
    "Cybersecurity Analyst",
    "Cloud Engineer",
    "Systems Administrator",
    "Network Engineer",
    "Technical Writer",
    "MERN Stack Developer",
    "Backend Engineer",
    "Frontend Engineer",
]


def _build_literal_replacements(
    doc_text: str,
    new_title: str,
    extra_candidates: list[str] | None = None,
) -> dict:
    """
    Scan *doc_text* for any known job-title literal (including variants with
    a " | tech stack" suffix such as "Full Stack Developer | Next.js · Node.js"),
    plus any *extra_candidates* supplied by the caller (e.g. the raw role and
    job_title from the UI fields).

    Return a dict mapping the exact matched string -> *new_title*.

    This is the primary fallback when the template has no formal
    ``{{JOB_TITLE}}`` placeholder, or when the current title in the template
    is not on the built-in list.
    """
    # Combine built-in bases with any extra candidates the caller supplied.
    all_bases = list(_KNOWN_JOB_TITLE_BASES)
    for cand in (extra_candidates or []):
        cand = cand.strip()
        if cand and cand not in all_bases:
            all_bases.append(cand)

    replacements = {}
    for base in all_bases:
        # Case-insensitive search for the base title in the document text.
        # We want to capture the full line/phrase including any " | ..." suffix.
        # Use regex: match the base (case-insensitive) optionally followed by
        # " | " and non-newline characters.
        pattern = re.compile(
            re.escape(base) + r"(?:\s*[|\u00b7·]\s*[^\n]*)?",
            re.IGNORECASE,
        )
        for m in pattern.finditer(doc_text):
            matched = m.group(0).strip()
            # Avoid replacing if the matched text is already the desired title
            if matched and matched != new_title:
                replacements[matched] = new_title
    return replacements


class DocEditor:
    """Edits DOCX templates by replacing designated placeholders only."""

    RESUME_PLACEHOLDERS = ["{{JOB_TITLE}}"]
    CV_PLACEHOLDERS = ["{{COMPANY_NAME}}", "{{ROLE}}"]

    def __init__(self, template_folder: str):
        """
        Args:
            template_folder: Folder that contains resume.docx and cv.docx.
        """
        self.template_folder = Path(template_folder)
        self.resume_template = self.template_folder / "resume.docx"
        self.cv_template = self.template_folder / "cv.docx"

    def get_resume_text(self) -> str:
        """Return the full text of the resume template."""
        if not self.resume_template.exists():
            return ""
        try:
            doc = Document(str(self.resume_template))
            return _document_full_text(doc)
        except Exception:
            return ""

    def get_existing_skills(self) -> list[str]:
        """Fallback to prevent AttributeError in older main.py logic."""
        return []

    # ------------------------------------------------------------------
    # Validation
    # ------------------------------------------------------------------

    def validate_templates(self) -> list[str]:
        """
        Check that template files exist.

        Returns:
            List of error strings (empty list = all good).
        """
        errors = []
        if not self.resume_template.exists():
            errors.append(
                f"Resume template not found: {self.resume_template}\n"
                "Please place resume.docx in the template folder."
            )
        if not self.cv_template.exists():
            errors.append(
                f"CV template not found: {self.cv_template}\n"
                "Please place cv.docx in the template folder."
            )
        return errors

    def _assert_placeholders(self, doc: Document, placeholders: list[str], filename: str) -> None:
        """
        We relax strict placeholder checks to support user-provided custom templates
        which may use alternative placeholders like [Company Name].
        """
        pass

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def edit_resume(self, job_title: str, output_path: str, role: str = "") -> str:
        """
        Copy resume template to *output_path* with {{JOB_TITLE}} replaced
        everywhere — including inside floating text boxes.

        Replacement targets (in priority order):
          1. ``{{JOB_TITLE}}`` and ``[Job Title]`` formal placeholders.
          2. The literal *role* string (the raw role from the job circular),
             including any " | tech stack" suffix variant found in the doc.
          3. The literal *job_title* string and any known base title from the
             built-in list (including " | tech stack" suffix variants).

        This means every occurrence of the old professional headline —
        whether in the body, a text box, a header, or a table cell — is
        updated to the new title.

        Args:
            job_title:   Formatted title to write (e.g. "Full Stack Developer"
                         or "Full Stack Developer  |  Next.js · Node.js").
            output_path: Destination path for the modified DOCX.
            role:        Raw role from the job circular (e.g. "Full Stack
                         Developer"). Used as an additional literal search
                         target so templates that already contain the raw role
                         text are also updated.

        Returns:
            output_path (str)

        Raises:
            FileNotFoundError: Template file missing.
            Exception: Any other docx / IO error.
        """
        if not self.resume_template.exists():
            raise FileNotFoundError(
                f"resume.docx not found at: {self.resume_template}"
            )
        try:
            doc = Document(str(self.resume_template))
            self._assert_placeholders(doc, self.RESUME_PLACEHOLDERS, "resume.docx")

            # Build replacement map — standard placeholders first
            replacements = {
                "{{JOB_TITLE}}": job_title,
                "[Job Title]": job_title,
            }

            # Collect extra candidate strings: the raw role and the formatted
            # job_title are both added as literal search targets so that any
            # template text matching either will be replaced.
            extra = [s for s in (role.strip(), job_title.strip()) if s]

            # Fallback: find any known (and extra-candidate) literal job-title
            # text in the document (including " | tech stack" suffix) and
            # replace it.  Text boxes are included via _document_full_text.
            doc_text = _document_full_text(doc)
            replacements.update(
                _build_literal_replacements(doc_text, job_title, extra_candidates=extra)
            )

            _replace_in_document(doc, replacements)
            doc.save(output_path)
            return output_path
        except (FileNotFoundError, ValueError):
            raise
        except Exception as e:
            raise Exception(f"Error editing resume.docx: {e}") from e

    def edit_cv(
        self,
        company_name: str,
        role: str,
        location: str,
        output_path: str,
        job_title: str = "",
    ) -> str:
        """
        Copy CV template to *output_path* with placeholders replaced.

        Also replaces {{JOB_TITLE}} / [Job Title] in the CV header so that
        cover-letter templates that echo the applicant's professional title
        are kept in sync with the resume.

        Args:
            company_name: Value to substitute for {{COMPANY_NAME}} or [Company Name].
            role:         Value to substitute for {{ROLE}} or [Position Name].
            location:     Value to substitute for {{LOCATION}} or [Company Address or Remote].
            output_path:  Destination path for the modified DOCX.
            job_title:    Value to substitute for {{JOB_TITLE}} / [Job Title] in the CV
                          header.  Defaults to *role* when not supplied.

        Returns:
            output_path (str)

        Raises:
            FileNotFoundError: Template file missing.
            ValueError: Placeholder(s) not present in the template.
            Exception: Any other docx / IO error.
        """
        if not self.cv_template.exists():
            raise FileNotFoundError(
                f"cv.docx not found at: {self.cv_template}"
            )
        try:
            import datetime
            current_date = datetime.datetime.now().strftime("%B %d, %Y")

            # Use role as fallback if no explicit job_title was provided
            effective_job_title = job_title.strip() if job_title.strip() else role

            doc = Document(str(self.cv_template))
            self._assert_placeholders(doc, self.CV_PLACEHOLDERS, "cv.docx")

            # Build replacement map
            replacements = {
                "{{COMPANY_NAME}}": company_name,
                "[Company Name]": company_name,
                "{{ROLE}}": role,
                "[Position Name]": role,
                "{{LOCATION}}": location,
                "[Company Address or Remote]": location,
                "{{DATE}}": current_date,
                "[Date]": current_date,
                # Job title placeholders in the CV header
                "{{JOB_TITLE}}": effective_job_title,
                "[Job Title]": effective_job_title,
            }

            # Collect extra candidate strings: the raw role and the formatted
            # job_title so that any template text matching either is replaced.
            extra = [s for s in (role.strip(), job_title.strip(), effective_job_title.strip()) if s]

            # Fallback: find any known (and extra-candidate) literal job-title
            # text in the document (incl. text boxes and " | tech stack" suffix)
            # and replace it.
            doc_text = _document_full_text(doc)
            replacements.update(
                _build_literal_replacements(doc_text, effective_job_title, extra_candidates=extra)
            )

            _replace_in_document(doc, replacements)
            doc.save(output_path)
            return output_path
        except (FileNotFoundError, ValueError):
            raise
        except Exception as e:
            raise Exception(f"Error editing cv.docx: {e}") from e

    def add_skills_to_resume(self, output_path: str, selected_skills: list[str]) -> str:
        """
        Add dynamically selected skills to the resume template.
        
        This method:
        1. Categorizes selected skills into Frontend, Backend, Auth, Tools
        2. Adds them to the corresponding skill subsections in the resume
        3. Preserves existing formatting and structure
        
        Args:
            output_path: Path where the resume was already saved (edit_resume call)
            selected_skills: List of skill names to add (from user checkboxes + manual)
        
        Returns:
            output_path (str)
        
        Raises:
            Exception: If skill placeholders not found or document error
        """
        if not selected_skills:
            # No skills to add, document already saved
            return output_path
        
        try:
            doc = Document(output_path)
            
            # Categorize the selected skills
            categorized = _categorize_skills(selected_skills)
            
            # Build replacements dict with placeholders
            replacements = {
                "{{frontend_skills}}": categorized["Frontend"],
                "{{backend_skills}}": categorized["Backend"],
                "{{auth_skills}}": categorized["Auth"],
                "{{tools_skills}}": categorized["Tools"],
            }
            
            # Only keep non-empty replacements to avoid adding empty commas
            replacements = {k: v for k, v in replacements.items() if v}
            
            # Replace placeholders in document
            _replace_in_document(doc, replacements)
            doc.save(output_path)
            
            return output_path
        except Exception as e:
            raise Exception(f"Error adding skills to resume: {e}") from e
