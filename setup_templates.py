"""
setup_templates.py
Generates sample resume.docx and cv.docx in the template folder
configured in config.json.

Run this once before first use:
    python setup_templates.py

The generated documents contain the required placeholders:
  resume.docx  →  {{JOB_TITLE}}
  cv.docx      →  {{COMPANY_NAME}}  and  {{ROLE}}

All surrounding content (skills, experience, education, etc.) is
real placeholder text so you can see how a finished document will look.
Replace this content with your actual information before building the EXE.
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def _set_run_style(run, bold=False, size_pt=11, colour: tuple | None = None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    if colour:
        run.font.color.rgb = RGBColor(*colour)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_style(run, bold=True, size_pt=13 if level == 1 else 11,
                   colour=(31, 73, 125) if level == 1 else (0, 0, 0))
    # Thin rule under level-1 headings
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run_style(run, size_pt=10)


def _normal(doc: Document, text: str, bold=False, size_pt=10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_style(run, bold=bold, size_pt=size_pt)


def _spacer(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


# ──────────────────────────────────────────────────────────────────────────────
# resume.docx
# ──────────────────────────────────────────────────────────────────────────────

def create_resume(output_path: Path) -> None:
    """
    Create a sample resume with the {{JOB_TITLE}} placeholder.
    The placeholder is used as the document's professional headline.
    """
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Name / contact header ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run("MD. FAYSAL AHMED")
    _set_run_style(name_run, bold=True, size_pt=18, colour=(31, 73, 125))

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(
        "Dhaka, Bangladesh  |  +880-1700-000000  |  faysal@example.com  |  linkedin.com/in/faysal"
    )
    _set_run_style(contact_run, size_pt=9, colour=(80, 80, 80))

    _spacer(doc)

    # ── {{JOB_TITLE}} placeholder — professional headline ──
    headline_para = doc.add_paragraph()
    headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline_run = headline_para.add_run("{{JOB_TITLE}}")
    _set_run_style(headline_run, bold=True, size_pt=12, colour=(200, 100, 0))

    _spacer(doc)

    # ── Professional Summary ──
    _heading(doc, "PROFESSIONAL SUMMARY")
    _normal(doc,
        "Results-driven technology professional with 5+ years of experience delivering "
        "scalable software solutions. Adept at cross-functional collaboration, problem-solving "
        "under pressure, and translating complex requirements into clean, maintainable code.",
        size_pt=10)

    _spacer(doc)

    # ── Skills ──
    _heading(doc, "TECHNICAL SKILLS")
    _bullet(doc, "Languages: Python, JavaScript (ES2022), TypeScript, SQL")
    _bullet(doc, "Frameworks: FastAPI, Django, React, Next.js, Node.js")
    _bullet(doc, "Databases: PostgreSQL, MySQL, MongoDB, Redis")
    _bullet(doc, "DevOps: Docker, GitHub Actions, Linux, AWS EC2/S3")
    _bullet(doc, "Tools: Git, VS Code, Postman, Jira, Figma")

    _spacer(doc)

    # ── Experience ──
    _heading(doc, "WORK EXPERIENCE")

    _normal(doc, "Senior Software Developer  |  Tech Solutions Ltd., Dhaka", bold=True, size_pt=11)
    _normal(doc, "March 2022 – Present", size_pt=9)
    _bullet(doc, "Designed and implemented RESTful APIs serving 500 k+ daily requests.")
    _bullet(doc, "Led migration of monolithic Django app to microservices, reducing deployment time by 40 %.")
    _bullet(doc, "Mentored 3 junior developers; introduced code-review culture and CI/CD pipelines.")

    _spacer(doc)

    _normal(doc, "Software Developer  |  Digital Agency BD, Dhaka", bold=True, size_pt=11)
    _normal(doc, "June 2019 – February 2022", size_pt=9)
    _bullet(doc, "Built full-stack web applications using React + Django REST Framework.")
    _bullet(doc, "Integrated third-party payment gateways (bKash, SSLCommerz) for e-commerce clients.")
    _bullet(doc, "Maintained 98 % uptime across 12 client deployments on AWS.")

    _spacer(doc)

    # ── Education ──
    _heading(doc, "EDUCATION")
    _normal(doc, "B.Sc. in Computer Science & Engineering", bold=True, size_pt=11)
    _normal(doc, "University of Dhaka  |  Graduated: 2019  |  CGPA: 3.72 / 4.00", size_pt=10)

    _spacer(doc)

    # ── Projects ──
    _heading(doc, "NOTABLE PROJECTS")
    _bullet(doc, "JobApplicationAI — Tkinter + Claude AI desktop tool for PDF generation (this app).")
    _bullet(doc, "EduPortal — LMS platform with video streaming for 10 k+ students.")
    _bullet(doc, "BazaarTrack — Real-time price comparison engine scraping 20 Bangladeshi e-shops.")

    _spacer(doc)

    # ── Languages ──
    _heading(doc, "LANGUAGES")
    _bullet(doc, "Bengali — Native")
    _bullet(doc, "English — Professional working proficiency")

    doc.save(str(output_path))
    print(f"  ✅  resume.docx created at {output_path}")


# ──────────────────────────────────────────────────────────────────────────────
# cv.docx
# ──────────────────────────────────────────────────────────────────────────────

def create_cv(output_path: Path) -> None:
    """
    Create a sample cover-letter CV with {{COMPANY_NAME}} and {{ROLE}} placeholders.
    """
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # ── Sender block ──
    _normal(doc, "Md. Faysal Ahmed", bold=True, size_pt=12)
    _normal(doc, "Dhaka, Bangladesh", size_pt=10)
    _normal(doc, "+880-1700-000000  |  faysal@example.com", size_pt=10)
    _spacer(doc)

    # ── Date ──
    from datetime import date
    _normal(doc, date.today().strftime("%B %d, %Y"), size_pt=10)
    _spacer(doc)

    # ── Recipient block — uses placeholders ──
    _normal(doc, "The Hiring Manager", bold=True, size_pt=10)

    company_para = doc.add_paragraph()
    company_run = company_para.add_run("{{COMPANY_NAME}}")
    _set_run_style(company_run, bold=True, size_pt=11, colour=(200, 100, 0))

    _normal(doc, "Bangladesh", size_pt=10)
    _spacer(doc)

    # ── Subject line ──
    subject_para = doc.add_paragraph()
    subject_para.add_run("Subject: Application for the Position of ").font.size = Pt(10)
    role_run = subject_para.add_run("{{ROLE}}")
    _set_run_style(role_run, bold=True, size_pt=11, colour=(200, 100, 0))
    _spacer(doc)

    # ── Salutation ──
    _normal(doc, "Dear Hiring Manager,", size_pt=10)
    _spacer(doc)

    # ── Body paragraphs ──
    _normal(doc,
        "I am writing to express my strong interest in the {{ROLE}} position at {{COMPANY_NAME}}. "
        "With over five years of hands-on experience in software development and a proven track "
        "record of delivering high-quality digital solutions, I am confident that my skills and "
        "enthusiasm align well with your team's objectives.",
        size_pt=10)
    _spacer(doc)

    _normal(doc,
        "In my current role at Tech Solutions Ltd., I have designed and maintained RESTful APIs "
        "handling more than 500,000 daily requests, led a successful microservices migration "
        "that cut deployment time by 40 %, and mentored junior colleagues through code reviews "
        "and technical workshops. These experiences have sharpened both my engineering skills "
        "and my ability to communicate clearly across teams.",
        size_pt=10)
    _spacer(doc)

    _normal(doc,
        "I am particularly excited about the opportunity at {{COMPANY_NAME}} because of your "
        "reputation for innovation and commitment to excellence. I believe my background in "
        "Python, modern web frameworks, and cloud infrastructure would allow me to contribute "
        "meaningfully from day one.",
        size_pt=10)
    _spacer(doc)

    _normal(doc,
        "I have attached my resume for your review and would welcome the opportunity to discuss "
        "how my experience can benefit {{COMPANY_NAME}}. Please feel free to contact me at "
        "+880-1700-000000 or faysal@example.com at your convenience.",
        size_pt=10)
    _spacer(doc)

    _normal(doc, "Thank you for your time and consideration.", size_pt=10)
    _spacer(doc)
    _spacer(doc)

    # ── Sign-off ──
    _normal(doc, "Sincerely,", size_pt=10)
    _spacer(doc)
    _normal(doc, "Md. Faysal Ahmed", bold=True, size_pt=11)

    doc.save(str(output_path))
    print(f"  ✅  cv.docx created at {output_path}")


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

def main() -> None:
    config_path = Path(__file__).parent / "config.json"

    # Load config to find the template folder
    try:
        with open(config_path, "r", encoding="utf-8") as fh:
            config = json.load(fh)
    except FileNotFoundError:
        print(f"ERROR: config.json not found at {config_path}")
        sys.exit(1)
    except json.JSONDecodeError as exc:
        print(f"ERROR: config.json is malformed — {exc}")
        sys.exit(1)

    template_folder_str = config.get("template_folder", "").strip()

    if not template_folder_str:
        # Fall back to templates/ subfolder next to this script
        template_folder = Path(__file__).parent / "templates"
        print(
            f"WARNING: 'template_folder' not set in config.json.\n"
            f"         Falling back to: {template_folder}"
        )
    else:
        template_folder = Path(template_folder_str)

    template_folder.mkdir(parents=True, exist_ok=True)
    print(f"\nCreating sample templates in: {template_folder}\n")

    resume_path = template_folder / "resume.docx"
    cv_path     = template_folder / "cv.docx"

    # Warn if files already exist (don't silently overwrite real resumes)
    for path in (resume_path, cv_path):
        if path.exists():
            answer = input(
                f"  ⚠  {path.name} already exists. Overwrite? [y/N] "
            ).strip().lower()
            if answer != "y":
                print(f"     Skipping {path.name}")
                if path == resume_path:
                    resume_path = None
                else:
                    cv_path = None

    if resume_path:
        create_resume(resume_path)
    if cv_path:
        create_cv(cv_path)

    print(
        "\nDone! Open resume.docx and cv.docx, replace the sample content\n"
        "with your real information, then run the main app.\n"
        "\nPlaceholders to keep intact:\n"
        "  resume.docx  →  {{JOB_TITLE}}\n"
        "  cv.docx      →  {{COMPANY_NAME}}  and  {{ROLE}}\n"
    )


if __name__ == "__main__":
    main()
